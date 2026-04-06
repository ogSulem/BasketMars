#!/usr/bin/env python3
"""
Создание diploma_reference.docx — шаблона Word с ГОСТ-стилями для КФУ.

Требования:
    pip install python-docx

Запуск:
    python3 create_reference.py

После запуска появится файл diploma_reference.docx.
Затем запустите:
    bash convert_diploma.sh

Параметры ГОСТ / методические указания КФУ (ВКР ПИ 2025):
    Шрифт основного текста:  Times New Roman, 14 pt
    Межстрочный интервал:    1.5
    Выравнивание текста:     по ширине
    Поля:  левое 30 мм, правое 15 мм, верхнее 20 мм, нижнее 20 мм
    Абзацный отступ:         1.25 см
    Заголовки:               TNR 14 pt обычный (без жирного), по центру, без отступа
    Подписи рисунков:        TNR 14 pt, по центру, обычный (без курсива — метод. указания КФУ)
    Код (листинг):           Courier New 10 pt, без отступа
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "diploma_reference.docx"


def set_page_margins(doc: Document) -> None:
    """Поля страницы по метод. указ. КФУ: левое 3 см, правое 1.5 см, верхнее 2 см, нижнее 2 см."""
    section = doc.sections[0]
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def set_run_font(run, name: str, size_pt: float, bold: bool = False,
                 italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    # Явно задаём шрифт для кириллицы (rFonts cs/eastAsia)
    rpr = run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)
    rfonts.set(qn("w:eastAsia"), name)
    rpr.insert(0, rfonts)


def configure_paragraph_format(para_fmt, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                first_line_cm: float = 1.25,
                                space_before_pt: float = 0,
                                space_after_pt: float = 0,
                                line_spacing_pt: float = None,
                                line_spacing_rule=None) -> None:
    from docx.enum.text import WD_LINE_SPACING
    para_fmt.alignment = align
    para_fmt.first_line_indent = Cm(first_line_cm)
    para_fmt.space_before = Pt(space_before_pt)
    para_fmt.space_after = Pt(space_after_pt)
    if line_spacing_pt is not None:
        para_fmt.line_spacing = Pt(line_spacing_pt)
        if line_spacing_rule is not None:
            para_fmt.line_spacing_rule = line_spacing_rule
    else:
        para_fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def apply_normal_style(doc: Document) -> None:
    """Основной текст: TNR 14pt, 1.5, по ширине, отступ 1.25 см."""
    from docx.enum.text import WD_LINE_SPACING
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    configure_paragraph_format(style.paragraph_format,
                               line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE)


def apply_heading_styles(doc: Document) -> None:
    """Заголовки 1–3: TNR 14pt обычный (без жирного), по центру, без отступа.
    Метод. указания КФУ: «без жирного выделения, без подчеркиваний»."""
    from docx.enum.text import WD_LINE_SPACING
    for level in range(1, 4):
        name = f"Heading {level}"
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style.font.bold = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12 if level == 1 else 6)
        pf.space_after = Pt(12)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.keep_with_next = True


def add_caption_style(doc: Document) -> None:
    """Подписи к рисункам/таблицам: TNR 14pt, по центру, без отступа.
    Метод. указания КФУ: «весь текст работы ... без курсива» — курсив запрещён."""
    from docx.enum.text import WD_LINE_SPACING
    try:
        style = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        style = doc.styles["Caption"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.italic = False
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(3)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_code_style(doc: Document) -> None:
    """Листинги кода: Courier New 10pt, без отступа, по левому краю."""
    from docx.enum.text import WD_LINE_SPACING
    try:
        style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        style = doc.styles["Code"]
    style.font.name = "Courier New"
    style.font.size = Pt(10)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_table_style(doc: Document) -> None:
    """Таблицы: TNR 12pt, без отступа, одинарный интервал."""
    from docx.enum.text import WD_LINE_SPACING
    try:
        style = doc.styles.add_style("Table Contents", WD_STYLE_TYPE.PARAGRAPH)
    except Exception:
        style = doc.styles["Table Contents"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_info_paragraph(doc: Document) -> None:
    """Добавляем пояснительный абзац в документ."""
    p = doc.add_paragraph()
    run = p.add_run(
        "Это reference-документ для pandoc. "
        "Он содержит стили ГОСТ/КФУ и не является содержательной частью диплома. "
        "Используйте: bash convert_diploma.sh"
    )
    set_run_font(run, "Times New Roman", 12, italic=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    doc = Document()
    set_page_margins(doc)
    apply_normal_style(doc)
    apply_heading_styles(doc)
    add_caption_style(doc)
    add_code_style(doc)
    add_table_style(doc)
    add_info_paragraph(doc)
    doc.save(OUTPUT)
    print(f"✓ Создан {OUTPUT} с ГОСТ-стилями (TNR 14pt, поля 30/15/20/20 мм, 1.5 интервал).")
    print(f"  Теперь запустите: bash convert_diploma.sh")


if __name__ == "__main__":
    main()
